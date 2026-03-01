"""Tests for Word document tools in mcp_server.py."""

import json
import os

import pytest
from docx import Document

from onit_office.mcp_server import (
    create_document,
    add_document_content,
    read_document,
    modify_document,
)


class TestCreateDocument:
    """Tests for create_document()."""

    def test_basic_creation(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "basic.docx"),
            title="My Document",
        ))
        assert result["status"] == "success"
        assert result["title"] == "My Document"
        assert os.path.exists(result["document_file"])

    def test_default_path(self, tmp_data_dir):
        result = json.loads(create_document(title="Default"))
        assert result["status"] == "success"
        assert os.path.exists(result["document_file"])

    def test_with_header(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "header.docx"),
            title="Header Doc",
            header_text="My Company",
        ))
        assert result["has_header"] is True

    def test_with_footer(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "footer.docx"),
            title="Footer Doc",
            footer_text="Page Footer",
        ))
        assert result["has_footer"] is True

    def test_with_logo(self, tmp_data_dir, sample_image):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "logo.docx"),
            title="Logo Doc",
            logo_path=sample_image,
            logo_width_inches=1.0,
        ))
        assert result["has_logo"] is True

    def test_with_header_and_logo(self, tmp_data_dir, sample_image):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "both.docx"),
            title="Both",
            header_text="Company",
            logo_path=sample_image,
        ))
        assert result["has_header"] is True
        assert result["has_logo"] is True

    def test_with_paragraph_content(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "para.docx"),
            title="Para Doc",
            content_type="paragraph",
            text="Initial paragraph content.",
        ))
        assert result["content_type"] == "paragraph"

    def test_with_heading_content(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "heading.docx"),
            content_type="heading",
            text="Section Heading",
        ))
        assert result["content_type"] == "heading"

    def test_with_bullets_content(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "bullets.docx"),
            title="Bullets",
            content_type="bullets",
            items=["First", "Second", "Third"],
        ))
        assert result["content_type"] == "bullets"

    def test_items_json_string(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "json_items.docx"),
            content_type="bullets",
            items='["A", "B"]',
        ))
        assert result["status"] == "success"

    def test_no_title(self, tmp_data_dir):
        result = json.loads(create_document(
            path=str(tmp_data_dir / "no_title.docx"),
        ))
        assert result["status"] == "success"
        assert result["title"] == ""


class TestAddDocumentContent:
    """Tests for add_document_content()."""

    def test_add_heading(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="heading",
            text="New Heading", level=2,
        ))
        assert result["status"] == "success"
        assert result["content_type"] == "heading"
        assert result["level"] == 2

    def test_heading_level_clamped(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="heading",
            text="Clamped", level=10,
        ))
        assert result["level"] == 4  # max is 4

    def test_heading_level_min(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="heading",
            text="Min", level=0,
        ))
        assert result["level"] == 1  # min is 1

    def test_add_paragraph(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="paragraph",
            text="A paragraph of text.",
        ))
        assert result["content_type"] == "paragraph"

    def test_paragraph_missing_text(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="paragraph",
        ))
        assert "error" in result

    def test_add_bullets(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="bullets",
            items=["Item 1", "Item 2", "Item 3"],
        ))
        assert result["item_count"] == 3

    def test_bullets_json_string(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="bullets",
            items='["X", "Y"]',
        ))
        assert result["item_count"] == 2

    def test_bullets_missing_items(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="bullets",
        ))
        assert "error" in result

    def test_add_image(self, docx_file, sample_image):
        result = json.loads(add_document_content(
            path=docx_file, content_type="image",
            image_path=sample_image,
        ))
        assert result["content_type"] == "image"

    def test_image_missing_path(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="image",
        ))
        assert "error" in result

    def test_image_nonexistent(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="image",
            image_path="/nonexistent/image.png",
        ))
        assert "error" in result

    def test_add_table(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="table",
            table_data=[["H1", "H2"], ["V1", "V2"]],
        ))
        assert result["table_size"] == "2x2"

    def test_table_json_string(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="table",
            table_data='[["A", "B"], ["1", "2"]]',
        ))
        assert result["table_size"] == "2x2"

    def test_table_missing_data(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="table",
        ))
        assert "error" in result

    def test_add_page_break(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="page_break",
        ))
        assert result["content_type"] == "page_break"

    def test_unknown_content_type(self, docx_file):
        result = json.loads(add_document_content(
            path=docx_file, content_type="unknown",
        ))
        assert "error" in result
        assert "valid_types" in result

    def test_nonexistent_file(self, tmp_data_dir):
        result = json.loads(add_document_content(
            path=str(tmp_data_dir / "nope.docx"),
            content_type="paragraph", text="Fail",
        ))
        assert "error" in result


class TestReadDocument:
    """Tests for read_document()."""

    def test_read_basic(self, docx_file):
        result = json.loads(read_document(path=docx_file))
        assert result["status"] == "success"
        assert result["paragraph_count"] >= 1
        assert result["section_count"] >= 1

    def test_read_with_content(self, docx_file):
        add_document_content(
            path=docx_file, content_type="paragraph", text="Test paragraph",
        )
        result = json.loads(read_document(path=docx_file))
        texts = [p["text"] for p in result["paragraphs"]]
        assert any("Test paragraph" in t for t in texts)

    def test_read_with_table(self, docx_file):
        add_document_content(
            path=docx_file, content_type="table",
            table_data=[["Name", "Value"], ["A", "1"]],
        )
        result = json.loads(read_document(path=docx_file))
        assert result["table_count"] == 1
        assert result["tables"][0]["rows"][0] == ["Name", "Value"]

    def test_read_with_header_footer(self, tmp_data_dir):
        path = str(tmp_data_dir / "hf.docx")
        create_document(
            path=path, title="HF",
            header_text="Header", footer_text="Footer",
        )
        result = json.loads(read_document(path=path))
        assert len(result["headers"]) > 0
        assert len(result["footers"]) > 0

    def test_nonexistent_file(self, tmp_data_dir):
        result = json.loads(read_document(
            path=str(tmp_data_dir / "nope.docx"),
        ))
        assert "error" in result


class TestModifyDocument:
    """Tests for modify_document()."""

    def test_update_paragraph_text(self, docx_file):
        # Add content to modify
        add_document_content(
            path=docx_file, content_type="paragraph", text="Original",
        )
        result = json.loads(modify_document(
            path=docx_file,
            updates=[{"paragraph_index": 1, "text": "Updated Text"}],
        ))
        assert result["status"] == "success"
        assert result["paragraphs_updated"] == 1

    def test_update_paragraph_style(self, docx_file):
        add_document_content(
            path=docx_file, content_type="paragraph", text="Style me",
        )
        result = json.loads(modify_document(
            path=docx_file,
            updates=[{"paragraph_index": 1, "style": "Heading 1"}],
        ))
        assert result["paragraphs_updated"] == 1

    def test_delete_paragraph(self, docx_file):
        add_document_content(
            path=docx_file, content_type="paragraph", text="Delete me",
        )
        # Get paragraph count
        read_result = json.loads(read_document(path=docx_file))
        para_count = read_result["paragraph_count"]

        result = json.loads(modify_document(
            path=docx_file, delete_indices=[para_count - 1],
        ))
        assert result["paragraphs_deleted"] == 1

    def test_json_string_updates(self, docx_file):
        add_document_content(
            path=docx_file, content_type="paragraph", text="JSON",
        )
        result = json.loads(modify_document(
            path=docx_file,
            updates='[{"paragraph_index": 1, "text": "JSON Updated"}]',
        ))
        assert result["paragraphs_updated"] == 1

    def test_json_string_deletes(self, docx_file):
        add_document_content(
            path=docx_file, content_type="paragraph", text="Del",
        )
        result = json.loads(modify_document(
            path=docx_file, delete_indices='[1]',
        ))
        assert result["paragraphs_deleted"] == 1

    def test_no_updates_or_deletes(self, docx_file):
        result = json.loads(modify_document(path=docx_file))
        assert "error" in result

    def test_invalid_paragraph_index(self, docx_file):
        result = json.loads(modify_document(
            path=docx_file,
            updates=[{"paragraph_index": 999, "text": "Fail"}],
        ))
        assert result["paragraphs_updated"] == 0

    def test_nonexistent_file(self, tmp_data_dir):
        result = json.loads(modify_document(
            path=str(tmp_data_dir / "nope.docx"),
            updates=[{"paragraph_index": 0, "text": "Fail"}],
        ))
        assert "error" in result
